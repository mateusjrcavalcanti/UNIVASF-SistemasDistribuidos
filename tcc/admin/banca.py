import os
from django import forms
from django.contrib import admin
from django.contrib.staticfiles import finders
from django.utils import timezone
from django.utils.html import format_html
from django.conf import settings
from django.db.models import Q, Count
from tcc.models import Banca, Horario, Trabalho
from django.contrib.auth.models import Group

from docx import Document


def getDocumentsUrl(trabalho):
    tccurl = f'{settings.MEDIA_URL}tcc/{trabalho.id}/'
    urls = {}

    if trabalho.banca.homologacao_coordenador and trabalho.banca.homologacao_orientador:
        declaracaoOrientador = f'Orientador_{trabalho.orientador.get_full_name()}.docx'
        urls.update({declaracaoOrientador: f'{tccurl}{declaracaoOrientador}'})

        for coorientador in trabalho.banca.coorientadores.all():
            declaracaoCoorientador = f'Coorientador_{coorientador.get_full_name()}.docx'
            urls.update(
                {declaracaoCoorientador: f'{tccurl}{declaracaoCoorientador}'})

        for avaliador in trabalho.banca.avaliadores.all():
            declaracaoAvaliador = f'Avaliador_{avaliador.get_full_name()}.docx'
            avaliacao = f'Avaliação_{avaliador.get_full_name()}.docx'
            avaliacaoExtra = f'AvaliaçãoExtra_{avaliador.get_full_name()}.docx'
            urls.update(
                {declaracaoAvaliador: f'{tccurl}{declaracaoAvaliador}',
                 avaliacao: f'{tccurl}{avaliacao}'})
            if trabalho.tipo == 'II':
                urls.update(
                    {avaliacaoExtra: f'{tccurl}{avaliacaoExtra}'})

    return urls


def generateDocuments(trabalho):

    tccdir = f'{settings.MEDIA_ROOT}/tcc/{trabalho.id}/'

    def applyFilter(document, Dictionary):
        for i in Dictionary:
            for p in document.paragraphs:
                if p.text.find(i) >= 0:
                    p.text = p.text.replace(i, Dictionary[i])

    if not os.path.exists(tccdir):
        os.makedirs(tccdir)

    Dictionary = {
        "::ORIENTADOR::": trabalho.orientador.get_full_name(),
        "::T_ORIENTADOR::": trabalho.orientador.titulo if trabalho.orientador.titulo is not None else '',
        "::TITULO::": f'{trabalho.titulo}',
        "::CANDIDATO::": trabalho.discente.get_full_name(),
        "::DATA::": timezone.localtime(trabalho.banca.horario.data).strftime('%d/%m/%Y') if trabalho.banca.horario.data is not None else '',
        "::HORA::": timezone.localtime(trabalho.banca.horario.data).strftime('%H:%M') if trabalho.banca.horario.data is not None else '',
        "::SEMESTRE::": f'{trabalho.semestre.ano}.{trabalho.semestre.numero}',
        "::CURSO::": f'{trabalho.semestre.curso.nome}',
        "::COORDENADOR::": trabalho.semestre.coordenador.get_full_name(),
        "::ESTAGIO::": "ESTAGIO?",
        "CECOMP": f'{trabalho.semestre.curso.sigla}',
                  "Prof. Dr. Jairson Barbosa Rodrigues": f'{trabalho.semestre.coordenador.titulo} {trabalho.semestre.coordenador.get_full_name()}',
    }

    declaracaoOrientador = Document(finders.find(
        'formularios/__Declaração Orientador - Modelo Canônico.docx'))
    applyFilter(declaracaoOrientador, Dictionary)
    declaracaoOrientador.save(
        f'{tccdir}Orientador_{trabalho.orientador.get_full_name()}.docx')

    for avaliador in trabalho.banca.avaliadores.all():
        Dictionary.update({"::T_AVAL::": f'{avaliador.titulo}',
                          "::AVAL::": avaliador.get_full_name()})

        declaracaoAvaliador = Document(finders.find(
            'formularios/__Declaração Avaliador - Modelo Canônico.docx'))
        applyFilter(declaracaoAvaliador, Dictionary)
        declaracaoAvaliador.save(
            f'{tccdir}Avaliador_{avaliador.get_full_name()}.docx')

        formularioAvaliador = Document(finders.find(
            'formularios/__Formulário de Avaliação TCC I - Modelo Canônico.docx' if trabalho.tipo == 'I' else 'formularios/__Formulário de Avaliação TCC II - Modelo Canônico.docx'))
        applyFilter(formularioAvaliador, Dictionary)
        formularioAvaliador.save(
            f'{tccdir}Avaliação_{avaliador.get_full_name()}.docx')

        if trabalho.tipo == 'II':
            formularioExtraAvaliador = Document(finders.find(
                'formularios/__Formulário de Avaliação TCC II (Complementar) - Modelo Canônico.docx'))
            applyFilter(formularioExtraAvaliador, Dictionary)
            formularioExtraAvaliador.save(
                f'{tccdir}AvaliaçãoExtra_{avaliador.get_full_name()}.docx')

    for coorientador in trabalho.banca.coorientadores.all():
        Dictionary.update({"::T_COORIENTADOR::": f'{coorientador.titulo}',
                          "::COORIENTADOR::": coorientador.get_full_name()})
        declaracaoCoorientador = Document(finders.find(
            'formularios/__Declaração Coorientador - Modelo Canônico.docx'))
        applyFilter(declaracaoCoorientador, Dictionary)
        declaracaoCoorientador.save(
            f'{tccdir}Coorientador_{coorientador.get_full_name()}.docx')


@admin.register(Banca)
class BancaAdmin(admin.ModelAdmin):
    list_per_page = 20
    list_max_show_all = 100
    save_as = True
    save_on_top = True

    actions_on_bottom = True

    list_display = ('trabalho', 'horario')
    readonly_fields = ('documentos',)

    def documentos(self, obj):
        html = ''
        if len(getDocumentsUrl(obj.trabalho)) > 0:
            for chave, valor in getDocumentsUrl(obj.trabalho).items():
                html += f'<a href="{valor}">{chave}</a><br>'
        return format_html(html)

    def get_queryset(self, request):
        qs = super(BancaAdmin, self).get_queryset(request)
        if request.user.is_superuser:
            return qs

        if request.user.groups.filter(name='Discentes').exists():
            return qs.filter(trabalho__discente__id=request.user.id)

        if qs.filter(Q(trabalho__semestre__coordenador__id=request.user.id) | Q(trabalho__discente_id=request.user.id) | Q(trabalho__orientador_id=request.user.id)).count() > 0:
            return qs.filter(Q(trabalho__semestre__coordenador__id=request.user.id) | Q(trabalho__discente_id=request.user.id) | Q(trabalho__orientador_id=request.user.id))
        else:
            return qs.filter(Q(trabalho__banca__avaliadores__id=request.user.id))

        # return qs.filter(Q(trabalho__semestre__coordenador__id=request.user.id) | Q(trabalho__discente_id=request.user.id) | Q(trabalho__orientador_id=request.user.id) | Q(trabalho__banca__avaliadores__id=request.user.id))
        # .annotate(count=Count('trabalho_id')).filter(count__lte=1)

    def has_change_permission(self, request, obj=None):
        return True

    def has_add_permission(self, request):
        if (request.user.is_superuser):
            return True
        if (Banca.objects.filter(trabalho__discente__id=request.user.id, trabalho__homologacao_orientador=True, trabalho__homologacao_coordenador=True).count() != Trabalho.objects.filter(discente__id=request.user.id).count()) or (Banca.objects.filter(trabalho__orientador__id=request.user.id, trabalho__homologacao_orientador=True, trabalho__homologacao_coordenador=True).count() != Trabalho.objects.filter(orientador__id=request.user.id).count()):
            return True
        return False

    def formfield_for_foreignkey(self, db_field, request, **kwargs):
        if db_field.name == "horario":
            if request.resolver_match.kwargs.get('object_id') is not None:
                kwargs["queryset"] = Horario.objects.filter(Q(banca=None) | Q(
                    banca__id=request.resolver_match.kwargs['object_id']))
            else:
                kwargs["queryset"] = Horario.objects.filter(banca=None)

        if db_field.name == "trabalho":
            if request.resolver_match.kwargs.get('object_id') is not None:
                kwargs["queryset"] = Trabalho.objects.filter(Q(discente__id=request.user.id) | Q(orientador__id=request.user.id) | Q(
                    banca__id=request.resolver_match.kwargs['object_id']), Q(homologacao_orientador=True), Q(homologacao_coordenador=True))
            else:
                kwargs["queryset"] = Trabalho.objects.filter(Q(discente__id=request.user.id) | Q(orientador__id=request.user.id), Q(
                    homologacao_orientador=True), Q(homologacao_coordenador=True))

        return super().formfield_for_foreignkey(db_field, request, **kwargs)

    def get_form(self, request, obj=None, **kwargs):

        form = super().get_form(request, obj=None, **kwargs)

        if obj is not None and obj.trabalho is not None and obj.trabalho.banca.horario is not None and obj.trabalho.homologacao_coordenador and obj.trabalho.homologacao_orientador:
            generateDocuments(obj.trabalho)

        if obj is None:
            # form.base_fields["avaliacoesresult"].disabled = True
            # form.base_fields["avaliacoesresult"].widget = forms.HiddenInput()

            # form.base_fields["resultado"].disabled = True
            # form.base_fields["resultado"].widget = forms.HiddenInput()

            form.base_fields["homologacao_coordenador"].disabled = True
            form.base_fields["homologacao_coordenador"].widget = forms.HiddenInput()

            form.base_fields["homologacao_orientador"].disabled = True
            form.base_fields["homologacao_orientador"].widget = forms.HiddenInput()
        else:
            # if (obj.trabalho.discente.id == request.user.id):
            #     form.base_fields["avaliacoesresult"].disabled = True
            #     form.base_fields["avaliacoesresult"].widget = forms.HiddenInput()
            if (obj.trabalho.orientador_id != request.user.id and obj.trabalho.discente.id != request.user.id and not request.user.is_superuser):
                form.base_fields["avaliadores"].disabled = True
                form.base_fields["coorientadores"].disabled = True
                form.base_fields["horario"].disabled = True
                form.base_fields["trabalho"].disabled = True
            if obj.trabalho.orientador_id != request.user.id:
                form.base_fields["homologacao_orientador"].disabled = True
                # form.base_fields["resultado"].disabled = True
            if obj.trabalho.semestre.coordenador.id != request.user.id:
                form.base_fields["homologacao_coordenador"].disabled = True

        return form
